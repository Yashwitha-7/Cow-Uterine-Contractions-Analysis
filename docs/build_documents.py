from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG_6263 = ROOT / "data" / "processed" / "cow_6263" / "figures"
FIG_6269 = ROOT / "data" / "processed" / "cow_6269" / "figures"

BROWN = RGBColor(91, 57, 37)
GREEN = RGBColor(49, 91, 58)
DARK = RGBColor(43, 39, 36)
MUTED = RGBColor(102, 94, 87)
CREAM = "F5EEE6"
LIGHT_GREEN = "E5EEE5"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


def set_font(run, size=11, bold=False, color=DARK, name="Arial", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc, title, running_label, preset="standard"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5 if preset == "compact" else 11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(5 if preset == "compact" else 6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BROWN, 14, 7),
        ("Heading 2", 13, GREEN, 11, 5),
        ("Heading 3", 11.5, DARK, 8, 3),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(running_label)
    set_font(r, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Hoffmann Laboratory  |  ")
    set_font(r, size=8, color=MUTED)
    r = p.add_run("Page ")
    set_font(r, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

    doc.core_properties.title = title
    doc.core_properties.subject = "Cow uterine contraction analysis"
    doc.core_properties.author = "Hoffmann Laboratory"


def add_cover(doc, title, subtitle, metadata, compact=False):
    doc.add_paragraph().paragraph_format.space_after = Pt(42 if compact else 88)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if compact else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HOFFMANN LABORATORY")
    set_font(r, size=10, bold=True, color=GREEN)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if compact else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, size=25 if compact else 28, bold=True, color=BROWN)
    p.paragraph_format.space_after = Pt(7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if compact else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_font(r, size=13, color=MUTED)
    p.paragraph_format.space_after = Pt(30)

    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if compact else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}: ")
        set_font(r, size=10.5, bold=True, color=DARK)
        r = p.add_run(value)
        set_font(r, size=10.5, color=DARK)
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p


def add_note(doc, label, text, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.65)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, bold=True, color=GREEN if fill == LIGHT_GREEN else BROWN)
    r = p.add_run(text)
    set_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, header_fill=CREAM):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        shade_cell(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_font(r, size=9.5, bold=True, color=BROWN)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, path, caption, width=6.55):
    path = Path(path)
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"Figure unavailable: {path.name}")
        set_font(r, size=9, italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", path.stem.replace("_", " ").title())
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    set_font(r, size=9, italic=True, color=MUTED)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.65)
    shade_cell(cell, "F7F7F7")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.strip().splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=8.8, color=DARK, name="Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_user_guide():
    doc = Document()
    configure_document(doc, "Cow Uterine Contraction Analysis — User Guide", "Repository User Guide", preset="compact")
    add_cover(
        doc,
        "Cow Uterine Contraction Analysis",
        "Simple user guide for installing the repository, processing data, and regenerating figures",
        [
            ("Repository", "github.com/Yashwitha-7/Cow-Uterine-Contractions-Analysis"),
            ("Audience", "Hoffmann Laboratory collaborators"),
            ("Prepared", "August 2026"),
        ],
        compact=True,
    )

    doc.add_heading("1. What this repository does", level=1)
    doc.add_paragraph(
        "The repository provides a local web application for ingesting cow contraction TXT files and optional bolus Excel files, checking data quality, reviewing possible strain-polarity reversals, detecting exploratory contraction candidates, calculating day/night and 24-hour statistics, and generating figures."
    )
    add_note(doc, "Scientific caution", "Detected peaks are candidate strain events. They are not confirmed physiological uterine contractions.")

    doc.add_heading("2. Before you begin", level=1)
    add_bullet(doc, "macOS, Linux, or Windows with a terminal application")
    add_bullet(doc, "Git")
    add_bullet(doc, "Python 3.12 recommended")
    add_bullet(doc, "Node.js 20 or later and npm")
    add_bullet(doc, "Raw contraction TXT files and, when available, the bolus Excel file")
    add_note(doc, "Data location", "Research data are deliberately excluded from GitHub. Obtain the raw data separately from the laboratory shared folder and never commit cow data to the public repository.", fill=CREAM)

    doc.add_heading("3. Download the repository", level=1)
    add_code(doc, "git clone https://github.com/Yashwitha-7/Cow-Uterine-Contractions-Analysis.git\ncd Cow-Uterine-Contractions-Analysis")

    doc.add_heading("4. Install the backend", level=1)
    add_code(doc, "cd backend\npython3.12 -m venv .venv\nsource .venv/bin/activate\npip install --upgrade pip\npip install -r requirements.txt")
    p = doc.add_paragraph("Windows activation command:")
    p.paragraph_format.keep_with_next = True
    add_code(doc, ".venv\\Scripts\\activate")

    doc.add_heading("5. Install the frontend", level=1)
    add_code(doc, "cd ../frontend\nnpm install")

    doc.add_heading("6. Start the application", level=1)
    doc.add_heading("Terminal 1 — backend", level=2)
    add_code(doc, "cd Cow-Uterine-Contractions-Analysis/backend\nsource .venv/bin/activate\nuvicorn app.main:app --reload")
    doc.add_heading("Terminal 2 — frontend", level=2)
    add_code(doc, "cd Cow-Uterine-Contractions-Analysis/frontend\nnpm run dev")
    doc.add_paragraph("Open the local address displayed by Vite, normally http://localhost:5173. The API documentation is available at http://127.0.0.1:8000/docs.")

    doc.add_page_break()
    doc.add_heading("7. Recommended analysis sequence", level=1)
    steps = [
        "Open Upload Data. Enter the cow ID and known calving date and time.",
        "Select Contractions TXT files and upload all TXT files for that cow in one batch.",
        "If available, upload one bolus Excel file for the same cow.",
        "Open QC Logs and Data Preview to confirm that files were read correctly.",
        "Open Polarity Review, enter the cow ID, and select Screen polarity.",
        "Review every flagged section using the raw strain, centered strain, movement channels, zero axis, time axis, and 30-minute context.",
        "Choose Keep polarity, Flip section, or Uncertain / exclude for every section.",
        "Use the hourly signal browser to inspect any hour manually.",
        "Open Phase 3 Processing and select Run Reviewed Analysis.",
        "Open Visual Analysis or Downloads to view the saved figures, statistics, and CSV files.",
    ]
    for step in steps:
        add_number(doc, step)

    doc.add_heading("8. How to make polarity decisions", level=1)
    add_table(
        doc,
        ["Decision", "Use when"],
        [
            ("Keep", "The flagged interval is consistent with surrounding signal direction."),
            ("Flip", "The complete flagged interval is clearly reversed relative to the surrounding recording."),
            ("Uncertain / exclude", "The signal direction cannot be established confidently."),
        ],
        [1.45, 5.05],
    )
    add_note(doc, "Important", "Only the reviewed interval is flipped. The 30-minute context is displayed for comparison and is not modified.")

    doc.add_heading("9. Movement interpretation", level=1)
    doc.add_paragraph("The contraction files contain a binary movement flag derived from the accelerometer and gyroscope channels:")
    add_bullet(doc, "0 means no movement detected.")
    add_bullet(doc, "5 means movement detected.")
    add_bullet(doc, "No intermediate values occurred in the available files.")
    doc.add_paragraph("The analysis also calculates acceleration and gyroscope variability independently. A strain peak near device-flagged or calculated movement is labeled movement-associated.")

    doc.add_heading("10. Where outputs are saved", level=1)
    add_code(doc, "data/\n  raw/cow_<id>/contractions/\n  raw/cow_<id>/bolus/\n  processed/cow_<id>/\n    quality_control/\n    statistics/\n    figures/\n    clocklab_exports/\n  database/hoffmann_lab.db")
    add_bullet(doc, "quality_control contains polarity screening, review manifests, decisions, and archived decisions.")
    add_bullet(doc, "statistics contains day/night CSV files and rhythm-summary JSON files.")
    add_bullet(doc, "figures contains regenerated PNG visualizations.")
    add_bullet(doc, "clocklab_exports contains CSV and AWD-formatted exports.")

    doc.add_heading("11. Reproducing the existing figures", level=1)
    doc.add_paragraph("To reproduce the same figures, collaborators need the same raw files, cow ID, calving timestamp, clock offsets, polarity decisions, software version, and analysis configuration. Recommended sharing package:")
    add_bullet(doc, "Raw contraction and bolus files in the laboratory shared folder")
    add_bullet(doc, "The relevant Git commit identifier")
    add_bullet(doc, "The polarity-decisions CSV")
    add_bullet(doc, "Cow metadata, including calving date and time")
    add_bullet(doc, "The generated statistics, figures, and processed CSV files")
    add_note(doc, "Reproducibility", "Do not compare regenerated outputs unless the raw files, review decisions, and code version are the same.")

    doc.add_heading("12. Verification commands", level=1)
    add_code(doc, "cd backend\n.venv/bin/python -m pytest tests -q\n\ncd ../frontend\nnpm run lint\nnpm run build")

    doc.add_page_break()
    doc.add_heading("13. Common problems", level=1)
    add_table(
        doc,
        ["Problem", "Recommended action"],
        [
            ("No module named app", "Run pytest from the backend folder, or use the documented command."),
            ("API returns Not Found", "Confirm the URL begins with /api and that the backend is running."),
            ("Analysis is locked", "Complete all pending polarity decisions first."),
            ("No figures appear", "Run Reviewed Analysis, then reload Visual Analysis or Downloads."),
            ("Duplicate upload blocked", "The cow already has that data type stored. Do not delete data unless the intended dataset is confirmed."),
            ("Port already in use", "Stop the earlier backend or frontend process before starting another."),
        ],
        [1.75, 4.75],
    )

    doc.add_heading("14. Safe data-handling rules", level=1)
    add_bullet(doc, "Do not modify files under data/raw.")
    add_bullet(doc, "Do not upload cow data, the SQLite database, or generated CSVs to the public GitHub repository.")
    add_bullet(doc, "Back up the data folder before resetting or replacing a cow dataset.")
    add_bullet(doc, "Record the code commit used for each sponsor-facing result package.")
    add_bullet(doc, "Describe clean peaks as contraction candidates, not confirmed contractions.")

    path = DOCS / "Cow_Contraction_Repository_User_Guide.docx"
    doc.save(path)
    return path


def build_report():
    doc = Document()
    configure_document(doc, "Cow Uterine Contraction Analysis — Detailed Report", "Analysis Report", preset="standard")
    add_cover(
        doc,
        "Cow Uterine Contraction Analysis",
        "Data-processing workflow, exploratory analyses, and current findings",
        [
            ("Laboratory", "Hoffmann Laboratory"),
            ("Data", "Cows 6263 and 6269"),
            ("Repository", "github.com/Yashwitha-7/Cow-Uterine-Contractions-Analysis"),
            ("Report date", "August 2026"),
        ],
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "A reproducible local workflow was developed to ingest, clean, review, analyze, and visualize cow uterine contraction sensor data. The current dataset contains contraction recordings from cows 6263 and 6269; cow 6263 also has bolus temperature and activity data. The workflow preserves all original observations, reconstructs timestamps for the current device, handles incomplete files and recording gaps, uses movement information to identify artifacts, and requires human review before applying possible polarity reversals."
    )
    doc.add_paragraph(
        "Across both cows, 203,013 contraction observations produced 1,438 detected strain peaks. After movement, polarity, and bad-signal exclusions, 26 peaks remained as clean contraction candidates. Cow 6263 showed a cluster of accepted events approximately two hours before calving. Cow 6269 did not show the same retained pattern, but a substantial amount of its activity remains excluded because of unresolved polarity. Bolus temperature for cow 6263 showed a clearer 24-hour rhythm than contraction-candidate activity."
    )
    add_note(doc, "Interpretation", "The retained events are signal-processing candidates, not physiologically confirmed uterine contractions. No predictive model was developed.")

    doc.add_heading("1. Study objective and scope", level=1)
    doc.add_paragraph("The objectives of this initial phase were to:")
    for text in (
        "Create a clean and traceable ingestion workflow for contraction TXT files and bolus Excel data.",
        "Reconstruct the contraction timeline while preserving incomplete files and true recording gaps.",
        "Identify movement, flat signal, and possible polarity problems before event analysis.",
        "Generate conservative contraction candidates for exploratory review.",
        "Compare contraction activity with bolus temperature on a common ten-minute timeline.",
        "Perform descriptive day/night and fixed 24-hour rhythm analyses.",
    ):
        add_bullet(doc, text)

    doc.add_heading("2. Available data", level=1)
    add_table(
        doc,
        ["Measurement", "Cow 6263", "Cow 6269"],
        [
            ("Contraction observations", "137,322", "65,691"),
            ("Contraction files", "85", "44"),
            ("Timeline span", "Approximately 90 hours", "Approximately 40 hours"),
            ("Bolus data", "Available", "Not available"),
            ("Calving time", "June 26, 2026, 3:54 AM", "July 4, 2026, 1:41 AM"),
        ],
        [2.15, 2.2, 2.2],
    )

    doc.add_heading("3. Sensor measurements", level=1)
    doc.add_heading("3.1 Contraction device", level=2)
    doc.add_paragraph("Each TXT row contains strain, three accelerometer axes, three gyroscope axes, a binary movement flag, and two unused user-input fields. The device developer confirmed that the movement flag is derived from the first six motion channels.")
    add_table(
        doc,
        ["Flag value", "Meaning", "Cow 6263 rows", "Cow 6269 rows"],
        [
            ("0", "No movement detected", "95,114", "40,693"),
            ("5", "Movement detected", "42,208", "24,998"),
        ],
        [1.0, 2.3, 1.6, 1.6],
        header_fill=LIGHT_GREEN,
    )
    doc.add_paragraph("No intermediate movement-flag values occurred in any of the 129 available contraction files. Approximately 30.7% of cow 6263 observations and 38.1% of cow 6269 observations were marked as movement.")

    doc.add_heading("3.2 Bolus device", level=2)
    doc.add_paragraph("Cow 6263 has 1,116 ten-minute bolus observations. Drink-cycle-corrected temperature was used for analysis, with raw temperature as a fallback. This avoids interpreting temporary cooling after water intake as a physiological temperature decrease. Raw temperature reached 34.66°C, whereas corrected temperature began at 38.27°C.")

    doc.add_heading("4. Processing challenges", level=1)
    add_bullet(doc, "The current contraction files do not contain a reliable per-observation time column.")
    add_bullet(doc, "The sampling interval varies and is approximately 1.8–2.2 seconds rather than exactly one second.")
    add_bullet(doc, "Hourly files can begin from different strain baselines.")
    add_bullet(doc, "The final file may be incomplete because the device was removed early.")
    add_bullet(doc, "Movement can produce strain peaks that resemble contraction events.")
    add_bullet(doc, "Possible polarity reversals cannot be established safely from skew alone.")
    add_bullet(doc, "No independent physiological reference confirms every contraction.")

    doc.add_page_break()
    doc.add_heading("5. Data-processing workflow", level=1)
    doc.add_heading("5.1 Ingestion and preservation", level=2)
    doc.add_paragraph("Every raw file is copied to an immutable cow-specific raw-data folder. Numeric sensor columns are validated and standardized, and all observations remain traceable to the source filename. Processed CSVs are saved separately from raw files.")

    doc.add_heading("5.2 Timestamp reconstruction", level=2)
    doc.add_paragraph("File start times are extracted from filenames. Reliable full-hour files provide sample-period estimates. Each row is assigned a timestamp from the file start, sample index, and estimated period. Partial files inherit a nearby reliable period and are shortened to the next file boundary when necessary to prevent timestamp overlap. True gaps create separate continuous recording segments.")

    doc.add_heading("5.3 Strain and movement quality control", level=2)
    doc.add_paragraph("Raw strain is retained. File-level median centering provides a robust QC representation that reduces arbitrary baseline differences without being dominated by isolated peaks. Acceleration and gyroscope magnitudes and their 30-second variability are calculated. A region is considered movement-associated when the binary device flag equals 5 or when calculated IMU variability is unusually high.")
    add_figure(doc, FIG_6263 / "cow_6263_daily_motion_sensor_rows.png", "Figure 1. Cow 6263 acceleration and gyroscope magnitude by day.")

    doc.add_heading("5.4 Polarity review", level=2)
    doc.add_paragraph("Files in which negative strain excursions dominate positive excursions are marked for human review. Consecutive flagged files are grouped. The reviewer sees the complete section, 30 minutes before and after it, exact time boundaries, strain, acceleration, and gyroscope signals. Decisions are Keep, Flip, or Uncertain / exclude. Raw files are never modified.")
    add_note(doc, "Current status", "Unresolved sections were temporarily marked uncertain and excluded from clean candidate counts. They remain available for review with the device developer.")

    doc.add_heading("5.5 Candidate detection", level=2)
    doc.add_paragraph("The reviewed centered strain is smoothed with a robust median window. Prominent peaks are detected using minimum prominence, temporal separation, and width rules. Each peak is classified as a clean candidate, movement-associated, uncertain-polarity, or bad-signal event. Clean candidates are exploratory and not confirmed contractions.")

    doc.add_heading("5.6 Ten-minute synchronization", level=2)
    doc.add_paragraph("The full-resolution contraction signal is preserved, while summary features and event counts are calculated in ten-minute intervals. Bolus observations are aligned with these intervals. Missing contraction coverage is explicitly recorded and not treated as zero activity.")

    doc.add_heading("5.7 Statistical analyses", level=2)
    add_bullet(doc, "Candidate count per valid recording hour")
    add_bullet(doc, "Day/night comparison using 6:00 AM–5:59 PM and 6:00 PM–5:59 AM")
    add_bullet(doc, "Fixed 24-hour cosinor summaries")
    add_bullet(doc, "Bolus temperature mean, median, variability, and daily rhythm")
    add_bullet(doc, "Candidate timing relative to documented calving")

    doc.add_page_break()
    doc.add_heading("6. Event-quality results", level=1)
    add_table(
        doc,
        ["Event category", "Cow 6263", "Cow 6269"],
        [
            ("Total detected peaks", "787", "651"),
            ("Movement-associated", "698", "413"),
            ("Uncertain polarity", "47", "210"),
            ("Bad signal", "23", "21"),
            ("Clean candidates", "19", "7"),
        ],
        [3.05, 1.75, 1.75],
    )
    doc.add_paragraph("Only 2.4% of cow 6263 peaks and 1.1% of cow 6269 peaks passed the current clean-candidate criteria. Movement was the dominant exclusion category. Polarity uncertainty had a particularly large effect on cow 6269.")

    doc.add_heading("7. Cow 6263 contraction findings", level=1)
    doc.add_paragraph("Cow 6263 produced 19 clean candidates. Nine occurred between approximately 1:45 AM and 2:03 AM on June 26, approximately 1 hour 51 minutes to 2 hours 9 minutes before the documented 3:54 AM calving time. Three clean candidates occurred within the final two hours, and five occurred two to six hours before calving.")
    add_figure(doc, FIG_6263 / "cow_6263_actogram_clean_candidate_peak_count.png", "Figure 2. Cow 6263 clean contraction-candidate count per ten-minute interval. White denotes missing samples.")
    doc.add_paragraph("The clustering is an important observation for this cow, but it cannot yet define a general prediction window. The all-candidate actogram is much denser, demonstrating how strongly movement and QC exclusions affect interpretation.")

    doc.add_heading("8. Cow 6269 contraction findings", level=1)
    doc.add_paragraph("Cow 6269 produced seven clean candidates. No clean candidate was retained in the final two hours before the documented 1:41 AM calving time. One occurred two to six hours before calving, and three occurred 12–24 hours before calving. Fifty-two peaks in the final two hours were excluded because they occurred in unresolved polarity sections.")
    add_figure(doc, FIG_6269 / "cow_6269_actogram_clean_candidate_peak_count.png", "Figure 3. Cow 6269 clean contraction-candidate count per ten-minute interval.")
    doc.add_paragraph("The absence of retained pre-calving candidates does not demonstrate absence of uterine contractions. Cow 6269 must be reassessed after polarity review because 210 detected peaks currently fall in uncertain-polarity regions.")

    doc.add_page_break()
    doc.add_heading("9. Cow 6263 bolus-temperature findings", level=1)
    add_table(
        doc,
        ["Statistic", "Result"],
        [
            ("Ten-minute observations", "1,116"),
            ("Rhythm-adjusted mean", "39.20°C"),
            ("Estimated 24-hour amplitude", "0.19°C"),
            ("Estimated peak phase", "Approximately 3:33 AM"),
            ("24-hour variance explained", "21.8%"),
            ("Daytime mean", "39.13°C"),
            ("Nighttime mean", "39.28°C"),
        ],
        [3.7, 2.8],
        header_fill=LIGHT_GREEN,
    )
    add_figure(doc, FIG_6263 / "cow_6263_bolus_temperature_actogram.png", "Figure 4. Cow 6263 drink-cycle-corrected bolus temperature by date and time of day.")
    doc.add_paragraph("Nighttime temperature was approximately 0.15°C warmer on average. The 24-hour temperature pattern was clearer than the contraction-candidate rhythm.")

    doc.add_heading("10. Parallel bolus and contraction view", level=1)
    doc.add_paragraph("Bolus temperature changes slowly, whereas strain candidates are short mechanical events. The shared ten-minute timeline permits parallel examination around calving without assuming that temperature and strain peaks should occur simultaneously.")
    add_figure(doc, FIG_6263 / "cow_6263_parallel_bolus_contraction_daily.png", "Figure 5. Cow 6263 bolus temperature and contraction strain on a common daily timeline.")

    doc.add_page_break()
    doc.add_heading("11. Day/night results", level=1)
    add_table(
        doc,
        ["Measurement", "Cow 6263", "Cow 6269"],
        [
            ("Day clean candidates", "10", "1"),
            ("Day valid hours", "36.17", "16.33"),
            ("Day candidates/hour", "0.276", "0.061"),
            ("Night clean candidates", "9", "6"),
            ("Night valid hours", "41.17", "24.00"),
            ("Night candidates/hour", "0.219", "0.250"),
        ],
        [3.05, 1.75, 1.75],
    )
    doc.add_paragraph("Cow 6263 had similar day and night rates. Cow 6269 had a higher nighttime rate. The cows do not show a consistent shared day/night pattern, and these comparisons are descriptive rather than population-level inference.")

    doc.add_heading("12. Circadian results", level=1)
    add_table(
        doc,
        ["Signal", "Estimated phase", "Variance explained"],
        [
            ("Cow 6263 candidates", "4:25 AM", "0.5%"),
            ("Cow 6269 candidates", "3:54 AM", "1.8%"),
            ("Cow 6263 bolus temperature", "3:33 AM", "21.8%"),
        ],
        [3.15, 1.65, 1.7],
    )
    doc.add_paragraph("The contraction-candidate rhythm was weak in both cows; therefore, the estimated phases should not be emphasized. Bolus temperature showed a clearer daily pattern.")

    doc.add_page_break()
    doc.add_heading("13. Main conclusions", level=1)
    for text in (
        "The workflow preserves raw data and produces traceable QC, processed data, statistics, and figures.",
        "Movement is a major component of the recorded strain signal and must be considered during peak interpretation.",
        "Cow 6263 showed a cluster of clean candidates approximately two hours before calving.",
        "Cow 6269 did not show the same retained pattern, but polarity uncertainty limits interpretation.",
        "Cow 6263 bolus temperature showed a measurable daily rhythm.",
        "The two cows show substantial individual and signal-quality differences.",
    ):
        add_bullet(doc, text)

    doc.add_heading("14. Limitations", level=1)
    add_bullet(doc, "Only two cows are available.")
    add_bullet(doc, "Only one cow has bolus data.")
    add_bullet(doc, "Current contraction timestamps are reconstructed.")
    add_bullet(doc, "Polarity review is incomplete.")
    add_bullet(doc, "No independent physiological reference confirms candidate contractions.")
    add_bullet(doc, "The meaning of the device’s movement algorithm is known at the binary-output level, but its internal threshold and time window remain undocumented.")
    add_bullet(doc, "No predictive or classification model is justified at the current sample size.")

    doc.add_heading("15. Recommended next steps", level=1)
    add_number(doc, "Review uncertain polarity sections with the device developer and rerun the analysis.")
    add_number(doc, "Use the next device version’s per-observation time column and document clock synchronization.")
    add_number(doc, "Collect additional cows with both contraction and bolus devices.")
    add_number(doc, "Place two contraction sensors on the same cow and compare agreement.")
    add_number(doc, "Standardize sensor placement, restart logging, calibration logging, and calving metadata.")
    add_number(doc, "Create an independently reviewed validation set before developing predictive models.")

    doc.add_heading("16. Reproducibility and sharing", level=1)
    doc.add_paragraph("The GitHub repository contains code and documentation but deliberately excludes raw cow data, processed results, figures, and the SQLite database. A complete shared-folder package should contain:")
    add_bullet(doc, "Raw contraction and bolus files")
    add_bullet(doc, "Cow metadata and calving timestamps")
    add_bullet(doc, "Polarity decisions")
    add_bullet(doc, "The Git commit identifier used for analysis")
    add_bullet(doc, "Generated processed CSVs, statistics, and figures")
    add_bullet(doc, "This report and the repository user guide")
    add_note(doc, "Repository", "https://github.com/Yashwitha-7/Cow-Uterine-Contractions-Analysis")

    path = DOCS / "Cow_Uterine_Contraction_Analysis_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    DOCS.mkdir(parents=True, exist_ok=True)
    print(build_user_guide())
    print(build_report())
